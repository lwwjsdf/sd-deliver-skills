#!/usr/bin/env python3
"""
design_performance_test.py — Generate a Performance Test Plan Word/Excel.

Usage:
    python3 design_performance_test.py \
      --dau 1000000 \
      --daily-events 5000000 \
      --retention-days 365 \
      --cloud AWS \
      --region ap-southeast-1 \
      --output-word ./references/performance-test-plan.docx \
      --output-excel ./references/performance-test-plan.xlsx

Dependencies:
    pip install python-docx openpyxl
"""

from __future__ import annotations

import argparse
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Missing dependency: python-docx")
    print("Run: ./venv/bin/pip install -r <skill-repo>/requirements.txt")
    sys.exit(1)

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
except ImportError:
    print("Missing dependency: openpyxl")
    print("Run: ./venv/bin/pip install -r <skill-repo>/requirements.txt")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Data
# ---------------------------------------------------------------------------


def _default_scenarios(dau: int, daily_events: int) -> List[Dict[str, str]]:
    return [
        {
            "No.": "PT-001",
            "Module": "CDP",
            "Scenario": "Real-time Import",
            "Data Preparation": f"{dau} baseline users; single event with PII + general attrs",
            "Test Steps": "Use JMeter incremental concurrency 10/20/30/1000/2000/3000, 10 min each",
            "Expected Metrics": "QPS ≥ 1000; CPU/Memory ≤ 80%; Disk I/O ≤ 70%",
        },
        {
            "No.": "PT-002",
            "Module": "CDP",
            "Scenario": "Batch Import",
            "Data Preparation": "10M events CSV, PGP encrypted, uploaded to SFTP",
            "Test Steps": "ETL downloads, decrypts, encrypts PII fields, imports via sync tool",
            "Expected Metrics": "≥ 1M records/hour; encryption + import time measured",
        },
        {
            "No.": "PT-003",
            "Module": "CDP",
            "Scenario": "Event Analysis - 7 Days",
            "Data Preparation": f"≥ {max(daily_events * 7, 200_000_000)} events",
            "Test Steps": "Query total count grouped by daily, no global filter, cache disabled",
            "Expected Metrics": "Response time ≤ 5s",
        },
        {
            "No.": "PT-004",
            "Module": "CDP",
            "Scenario": "Event Analysis - 30 Days",
            "Data Preparation": f"≥ {max(daily_events * 30, 200_000_000)} events",
            "Test Steps": "Query total count grouped by daily, no global filter, cache disabled",
            "Expected Metrics": "Response time ≤ 5s",
        },
        {
            "No.": "PT-005",
            "Module": "CDP",
            "Scenario": "Funnel Analysis - 7 Days",
            "Data Preparation": f"≥ {max(daily_events * 7, 200_000_000)} events",
            "Test Steps": "3-step funnel, 7-day window, no filter, cache disabled",
            "Expected Metrics": "Response time ≤ 5s",
        },
        {
            "No.": "PT-006",
            "Module": "MAE",
            "Scenario": "Canvas Execution",
            "Data Preparation": "Prepared users/events from above scenarios",
            "Test Steps": "Publish one-time Canvas; measure startup to sync completion",
            "Expected Metrics": "Execution ≤ 1min; rule accuracy ≥ 99%; sync accuracy ≥ 99.9%",
        },
        {
            "No.": "PT-007",
            "Module": "MAE",
            "Scenario": "Journey Email Sending",
            "Data Preparation": "Prepared users/events",
            "Test Steps": "JMeter trigger 1000/2000/3000/6000 events; count successful emails/sec",
            "Expected Metrics": "≥ 1000 emails/minute; CPU/Memory ≤ 80%; Disk I/O ≤ 70%",
        },
    ]


# ---------------------------------------------------------------------------
# Word generation
# ---------------------------------------------------------------------------


def _add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.color.rgb = RGBColor(0x04, 0x5C, 0x43)
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    return heading


def _add_table(doc: Document, rows: List[Dict[str, str]]):
    if not rows:
        return
    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        shading_elm = hdr_cells[i]._tc.get_or_add_tcPr()
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
        shading_elm.append(parse_xml(r'<w:shd {} w:fill="00A870"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')))
    for row in rows:
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row.get(h, ""))


def generate_word_report(
    output_path: str,
    dau: int,
    daily_events: int,
    retention_days: int,
    cloud: str,
    region: str,
    include_cdp: bool,
    include_ma: bool,
):
    doc = Document()
    title = doc.add_heading("Performance Test Plan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x04, 0x5C, 0x43)

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()

    _add_heading(doc, "1. Introduction / Purpose", level=1)
    doc.add_paragraph(
        "Define the overall plan, execution strategy and acceptance criteria for performance testing "
        "of the CDP and MAE systems in the UAT environment."
    )

    _add_heading(doc, "2. Objective", level=1)
    doc.add_paragraph(
        "Verify system performance under expected production-scale load and quantify the impact of "
        "PII data encryption on core components."
    )

    _add_heading(doc, "3. Scope of Testing", level=1)
    _add_table(doc, _default_scenarios(dau, daily_events))

    _add_heading(doc, "4. Expected Outcomes", level=1)
    doc.add_paragraph("- Complete Performance Test Report")
    doc.add_paragraph("- Clear go-live readiness conclusion")
    doc.add_paragraph("- Quantified encryption overhead")

    _add_heading(doc, "5. Testing Environment", level=1)
    env_rows = [
        {"Item": "Cloud", "Value": cloud},
        {"Item": "Region", "Value": region},
        {"Item": "DAU", "Value": f"{dau:,}"},
        {"Item": "Daily Events", "Value": f"{daily_events:,}"},
        {"Item": "Retention Days", "Value": str(retention_days)},
        {"Item": "Products", "Value": ", ".join(["CDP" if include_cdp else "", "MA" if include_ma else ""]).strip(", ")},
    ]
    _add_table(doc, env_rows)

    _add_heading(doc, "6. Test Assumptions and Risks", level=1)
    doc.add_paragraph("- UAT environment matches production")
    doc.add_paragraph("- Test data distribution matches production")
    doc.add_paragraph("- UAT resources are exclusively used during testing")

    _add_heading(doc, "7. Test Execution", level=1)
    _add_heading(doc, "7.1 Entry Criteria", level=2)
    doc.add_paragraph("Platform deployed, test data ready, load generators ready, plan approved.")
    _add_heading(doc, "7.2 Exit Criteria", level=2)
    doc.add_paragraph("All scenarios executed; metrics meet requirements or boundary determined; report submitted.")

    _add_heading(doc, "8. Roles and Responsibilities", level=1)
    role_rows = [
        {"Role": "SD Application Team", "Responsibilities": "Plan, scripts, execution, analysis, report"},
        {"Role": "WKCDA System/Network Team", "Responsibilities": "System and network support"},
        {"Role": "WKCDA Support Team", "Responsibilities": "Review plan and results"},
    ]
    _add_table(doc, role_rows)

    _add_heading(doc, "9. Schedule", level=1)
    schedule_rows = [
        {"Activity": "Submit Performance Test Plan", "Responsible": "SD", "Start": "", "End": ""},
        {"Activity": "Review Plan", "Responsible": "WKCDA", "Start": "", "End": ""},
        {"Activity": "Execute Testing", "Responsible": "SD", "Start": "", "End": ""},
        {"Activity": "Submit Report", "Responsible": "SD", "Start": "", "End": ""},
    ]
    _add_table(doc, schedule_rows)

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Excel generation
# ---------------------------------------------------------------------------


def _write_excel_sheet(ws, rows: List[Dict[str, str]]):
    if not rows:
        ws.append(["No data"])
        return
    headers = list(rows[0].keys())
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="00A870", end_color="00A870", fill_type="solid")
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    for row in rows:
        ws.append([row.get(h, "") for h in headers])
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except Exception:
                pass
        ws.column_dimensions[col_letter].width = min(max(max_length + 2, 12), 60)


def generate_excel_report(
    output_path: str,
    dau: int,
    daily_events: int,
    retention_days: int,
    cloud: str,
    region: str,
    include_cdp: bool,
    include_ma: bool,
):
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "Scope"
    _write_excel_sheet(ws1, _default_scenarios(dau, daily_events))

    ws2 = wb.create_sheet("Environment")
    env_rows = [
        {"Item": "Cloud", "Value": cloud},
        {"Item": "Region", "Value": region},
        {"Item": "DAU", "Value": f"{dau:,}"},
        {"Item": "Daily Events", "Value": f"{daily_events:,}"},
        {"Item": "Retention Days", "Value": str(retention_days)},
        {"Item": "Products", "Value": ", ".join(["CDP" if include_cdp else "", "MA" if include_ma else ""]).strip(", ")},
    ]
    _write_excel_sheet(ws2, env_rows)

    ws3 = wb.create_sheet("Assumptions")
    _write_excel_sheet(ws3, [
        {"Assumption": "UAT environment matches production"},
        {"Assumption": "Test data distribution matches production"},
        {"Assumption": "UAT resources exclusively used during testing"},
    ])

    ws4 = wb.create_sheet("Schedule")
    _write_excel_sheet(ws4, [
        {"Activity": "Submit Performance Test Plan", "Responsible": "SD", "Start": "", "End": ""},
        {"Activity": "Review Plan", "Responsible": "WKCDA", "Start": "", "End": ""},
        {"Activity": "Execute Testing", "Responsible": "SD", "Start": "", "End": ""},
        {"Activity": "Submit Report", "Responsible": "SD", "Start": "", "End": ""},
    ])

    wb.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(description="Generate Performance Test Plan Word/Excel")
    parser.add_argument("--dau", type=int, required=True, help="Daily active users")
    parser.add_argument("--daily-events", type=int, required=True, help="Daily event volume")
    parser.add_argument("--retention-days", type=int, default=365, help="Data retention days")
    parser.add_argument("--cloud", default="AWS", help="Cloud vendor")
    parser.add_argument("--region", default="ap-southeast-1", help="Cloud region")
    parser.add_argument("--include-cdp", action="store_true", default=True, help="Include CDP")
    parser.add_argument("--include-ma", action="store_true", default=True, help="Include MA")
    parser.add_argument("--output-word", help="Output Word path")
    parser.add_argument("--output-excel", help="Output Excel path")
    args = parser.parse_args()

    if args.output_word:
        path = generate_word_report(
            args.output_word,
            args.dau,
            args.daily_events,
            args.retention_days,
            args.cloud,
            args.region,
            args.include_cdp,
            args.include_ma,
        )
        print(f"Word report generated: {path}")

    if args.output_excel:
        path = generate_excel_report(
            args.output_excel,
            args.dau,
            args.daily_events,
            args.retention_days,
            args.cloud,
            args.region,
            args.include_cdp,
            args.include_ma,
        )
        print(f"Excel report generated: {path}")


if __name__ == "__main__":
    main()
