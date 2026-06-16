#!/usr/bin/env python3
"""
validate_performance_test_plan.py — Validate generated Performance Test Plan Word/Excel.

Usage:
    ./venv/bin/python <skill-repo>/sd-infra/scripts/validate_performance_test_plan.py \
      --format word \
      --input ./references/performance-test-plan.docx

    ./venv/bin/python <skill-repo>/sd-infra/scripts/validate_performance_test_plan.py \
      --format excel \
      --input ./references/performance-test-plan.xlsx
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import List

REQUIRED_SECTIONS = [
    "Introduction",
    "Objective",
    "Scope",
    "Expected Outcomes",
    "Testing Environment",
    "Assumptions",
    "Test Execution",
    "Roles",
    "Schedule",
]

REQUIRED_SCOPE_COLUMNS = [
    "No.",
    "Module",
    "Scenario",
    "Data Preparation",
    "Test Steps",
    "Expected Metrics",
]


def _fail(messages: List[str]) -> int:
    print("❌ Performance Test Plan validation failed:")
    for msg in messages:
        print(f"  - {msg}")
    return 1


def validate_word(path: str) -> int:
    try:
        from docx import Document
    except ImportError:
        print("Missing dependency: python-docx")
        print("Run: ./venv/bin/pip install -r <skill-repo>/requirements.txt")
        return 1

    doc = Document(path)
    paragraphs_text = "\n".join(p.text for p in doc.paragraphs)

    # Collect text from tables as well
    table_texts: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_texts.append(cell.text)
    tables_text = "\n".join(table_texts)
    full_text = paragraphs_text + "\n" + tables_text

    errors: List[str] = []
    for section in REQUIRED_SECTIONS:
        if section not in full_text:
            errors.append(f"缺少章节: {section}")

    # Scope table should contain scenario rows
    if "Scenario" not in full_text or "Expected Metrics" not in full_text:
        errors.append("Scope 表缺少关键列（Scenario / Expected Metrics）")

    if not any(line.startswith("PT-") for line in full_text.splitlines()):
        errors.append("Scope 中未找到场景编号（PT-001 等）")

    if errors:
        return _fail(errors)

    print("✅ Word Performance Test Plan is valid.")
    return 0


def validate_excel(path: str) -> int:
    try:
        import openpyxl
    except ImportError:
        print("Missing dependency: openpyxl")
        print("Run: ./venv/bin/pip install -r <skill-repo>/requirements.txt")
        return 1

    wb = openpyxl.load_workbook(path, data_only=True)
    errors: List[str] = []

    if "Scope" not in wb.sheetnames:
        errors.append(f"缺少 sheet 'Scope'，实际 sheets: {wb.sheetnames}")
    else:
        ws = wb["Scope"]
        if ws.max_row < 2:
            errors.append("Scope sheet 没有数据行")
        else:
            headers = [
                str(cell.value).strip() if cell.value is not None else ""
                for cell in ws[1]
            ]
            headers = [h for h in headers if h]
            missing = [h for h in REQUIRED_SCOPE_COLUMNS if h not in headers]
            if missing:
                errors.append(f"Scope 列缺失: {missing}")

            scenario_count = 0
            for row in ws.iter_rows(min_row=2, values_only=True):
                if row and any(cell is not None for cell in row):
                    scenario_count += 1
            if scenario_count == 0:
                errors.append("Scope sheet 没有场景数据")

    for sheet in ["Environment", "Assumptions", "Schedule"]:
        if sheet not in wb.sheetnames:
            errors.append(f"缺少 sheet '{sheet}'")

    if errors:
        return _fail(errors)

    print("✅ Excel Performance Test Plan is valid.")
    return 0


def main():
    parser = argparse.ArgumentParser(description="Validate Performance Test Plan")
    parser.add_argument("--format", choices=["word", "excel"], required=True)
    parser.add_argument("--input", required=True, help="Path to Word or Excel plan")
    args = parser.parse_args()

    if not Path(args.input).exists():
        print(f"❌ File not found: {args.input}")
        sys.exit(1)

    if args.format == "word":
        rc = validate_word(args.input)
    else:
        rc = validate_excel(args.input)

    sys.exit(rc)


if __name__ == "__main__":
    main()
