#!/usr/bin/env python3
"""
Markdown to Word Converter
支持中文字体和格式保留

用法：
  python3 md2docx.py input.md output.docx
"""

import sys
import subprocess
import tempfile
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 字体配置 ──────────────────────────────────────────────────────────────────
FONT_CONFIG = {
    # 西文
    'heading_west': 'Arial',
    'body_west': 'Times New Roman',
    'code_west': 'Courier New',
    # 中文
    'heading_east': 'SimHei',      # 黑体
    'body_east': 'SimSun',         # 宋体
    'code_east': 'Courier New',
}

STYLE_FONTS = {
    'Title':            ('heading_west', 'heading_east', 18, True),
    'Heading 1':        ('heading_west', 'heading_east', 16, True),
    'Heading 2':        ('heading_west', 'heading_east', 14, True),
    'Heading 3':        ('heading_west', 'heading_east', 12, True),
    'Heading 4':        ('heading_west', 'heading_east', 11, True),
    'Heading 5':        ('heading_west', 'heading_east', 11, False),
    'Heading 6':        ('heading_west', 'heading_east', 11, False),
    'Normal':           ('body_west', 'body_east', 11, False),
    'Table Normal':     ('body_west', 'body_east', 10, False),
    'Code':             ('code_west', 'code_east', 10, False),
}


def set_run_font(run, west_font, east_font, size_pt, bold=False, color=None):
    """设置 run 的字体（支持中西文混排）"""
    font = run.font
    font.name = west_font
    font.size = Pt(size_pt)
    font.bold = bold
    if color:
        font.color.rgb = color
    
    # 设置中文字体（通过 XML）
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    
    # 设置东亚字体
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    
    rFonts.set(qn('w:eastAsia'), east_font)
    rFonts.set(qn('w:ascii'), west_font)
    rFonts.set(qn('w:hAnsi'), west_font)


def format_paragraph(para, style_name):
    """格式化段落中的所有 run"""
    if style_name not in STYLE_FONTS:
        style_name = 'Normal'
    
    west_font, east_font, size_pt, bold = STYLE_FONTS[style_name]
    
    for run in para.runs:
        set_run_font(run, west_font, east_font, size_pt, bold)
    
    # 设置段落对齐
    if style_name == 'Title':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def format_table(table):
    """格式化表格"""
    west_font, east_font, size_pt, bold = STYLE_FONTS['Table Normal']
    
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, west_font, east_font, size_pt, bold)
                # 表格内容默认左对齐
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def format_code_blocks(doc):
    """格式化代码块（通常以特定样式或背景色标识）"""
    west_font, east_font, size_pt, bold = STYLE_FONTS['Code']
    
    for para in doc.paragraphs:
        # 检测代码块（通过样式或内容特征）
        text = para.text.strip()
        if text.startswith('```') or text.startswith('    ') or text.startswith('\t'):
            for run in para.runs:
                set_run_font(run, west_font, east_font, size_pt, bold)
                # 灰色背景
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def highlight_tbc_items(doc):
    """高亮所有 TBC 项目（黄色背景）"""
    import re
    
    YELLOW = RGBColor(0xFF, 0xFF, 0x00)
    
    def apply_yellow_highlight(run):
        """为 run 应用黄色高亮背景"""
        # 使用 shading 设置背景色
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        
        # 移除已有的 shading
        existing_shd = rPr.find(qn('w:shd'))
        if existing_shd is not None:
            rPr.remove(existing_shd)
        
        # 添加黄色背景
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'FFFF00')
        rPr.append(shd)
    
    # 处理段落中的 TBC
    for para in doc.paragraphs:
        for run in para.runs:
            if '[TBC' in run.text or 'TBC]' in run.text:
                apply_yellow_highlight(run)
    
    # 处理表格中的 TBC
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if '[TBC' in run.text or 'TBC]' in run.text:
                            apply_yellow_highlight(run)


def process_document(docx_path):
    """处理文档，应用字体和格式"""
    doc = Document(docx_path)
    
    # 处理段落
    for para in doc.paragraphs:
        style_name = para.style.name
        format_paragraph(para, style_name)
    
    # 处理表格
    for table in doc.tables:
        format_table(table)
    
    # 处理代码块
    format_code_blocks(doc)
    
    # 高亮 TBC 项目
    highlight_tbc_items(doc)
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    
    doc.save(docx_path)
    print(f"✓ Formatted: {docx_path}")


def md2docx(md_path, docx_path):
    """Markdown to Word 转换主函数"""
    md_path = Path(md_path)
    docx_path = Path(docx_path)
    
    if not md_path.exists():
        print(f"✗ Error: File not found: {md_path}")
        sys.exit(1)
    
    # 确保输出目录存在
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Step 1: pandoc 基础转换
    print(f"Converting {md_path.name} → {docx_path.name}...")
    
    cmd = [
        'pandoc',
        str(md_path),
        '-o', str(docx_path),
        '--from', 'markdown',
        '--to', 'docx',
        '--wrap', 'none',           # 不自动换行
        '--table-of-contents',      # 添加目录
        '--toc-depth', '3',         # 目录深度到 H3
    ]
    
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"✗ Pandoc error: {result.stderr}")
        sys.exit(1)
    
    print(f"✓ Pandoc conversion complete")
    
    # Step 2: python-docx 格式美化
    print("Applying Chinese font formatting...")
    process_document(docx_path)
    
    print(f"\n✓ Done: {docx_path}")
    print(f"  File size: {docx_path.stat().st_size / 1024:.1f} KB")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python3 md2docx.py <input.md> <output.docx>")
        print("\nExample:")
        print("  python3 md2docx.py TechSpec.md TechSpec.docx")
        sys.exit(1)
    
    md2docx(sys.argv[1], sys.argv[2])
