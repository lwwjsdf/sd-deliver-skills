"""
docx_to_yaml_skeleton.py
========================
一次性工具脚本：从 UAT_TestDataset_BusinessLogic_v1.docx 提取表格数据，
生成 rules/special/westk/business_logic.yaml 骨架。

运行方式：
    cd tracking-setup-e2e
    python3 scripts/docx_to_yaml_skeleton.py

依赖（可选）：
    pip install python-docx pyyaml

如果 docx 解析失败，脚本会直接使用硬编码数据生成 YAML。
"""

import os
import sys
import yaml

# ---------------------------------------------------------------------------
# 路径配置
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)          # tracking-setup-e2e/
DOCX_PATH = os.path.join(BASE_DIR, "refrences", "UAT_TestDataset_BusinessLogic_v1.docx")
OUTPUT_DIR = os.path.join(BASE_DIR, "rules", "special", "westk")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "business_logic.yaml")


# ---------------------------------------------------------------------------
# 尝试解析 docx（失败时静默降级到硬编码）
# ---------------------------------------------------------------------------
def try_parse_docx(path):
    """尝试用 python-docx 读取文档，返回 (tables, paragraphs) 或 (None, None)。"""
    try:
        from docx import Document
        doc = Document(path)
        return doc.tables, doc.paragraphs
    except Exception as e:
        print(f"[warn] docx 解析失败，使用硬编码数据：{e}", file=sys.stderr)
        return None, None


# ---------------------------------------------------------------------------
# 硬编码数据（正确性优先）
# ---------------------------------------------------------------------------
def build_hardcoded_data():
    data = {}

    # 1. meta
    data["meta"] = {
        "project": "westk",
        "version": "1.1",
        "platforms": ["MP", "Web"],
    }

    # 2. region_distribution（全局）
    data["region_distribution"] = {
        "mainland": 0.60,
        "hongkong": 0.30,
        "overseas": 0.10,
    }

    # 3. user_segments（Table 1）
    data["user_segments"] = {
        "L0": {
            "ratio": 0.60,
            "identities": ["unionid", "cookie_id"],
            "has_registration": False,
            "has_purchase": False,
            "has_membership": False,
            "membership_activated": False,
        },
        "L1": {
            "ratio": 0.20,
            "identities": ["mobile_or_email", "unionid", "cookie_id"],
            "has_registration": True,
            "has_purchase": False,
            "has_membership": False,
            "membership_activated": False,
        },
        "L2": {
            "ratio": 0.10,
            "identities": ["mobile_or_email", "unionid", "cookie_id"],
            "has_registration": True,
            "has_purchase": True,
            "has_membership": False,
            "membership_activated": False,
        },
        "L3": {
            "ratio": 0.03,
            "identities": ["mobile_or_email", "unionid", "cookie_id"],
            "has_registration": True,
            "has_purchase": True,
            "has_membership": True,
            "membership_activated": False,
        },
        "L4": {
            "ratio": 0.07,
            "identities": ["mobile_or_email", "unionid", "cookie_id"],
            "has_registration": True,
            "has_purchase": True,
            "has_membership": True,
            "membership_activated": True,
        },
    }

    # 4. identity_priority
    data["identity_priority"] = {
        "crm_master_id": {
            "priority": 0,
            "display": "CRM Master Customer ID",
            "sa_key": "$identity_login_id",
        },
        "email": {
            "priority": 1,
            "display": "Email",
            "sa_key": "$identity_email",
        },
        "mobile": {
            "priority": 1,
            "display": "Mobile",
            "sa_key": "$identity_mobile",
        },
        "unionid": {
            "priority": 2,
            "display": "WeChat UnionID",
            "sa_key": "$identity_unionid",
        },
        "cookie_id": {
            "priority": 2,
            "display": "CookieID",
            "sa_key": "$identity_anonymous_id",
        },
    }

    # 5. event_sequences（骨架）
    data["event_sequences"] = [
        {
            "name": "user_lifecycle",
            "condition": None,
            "events": [],
        },
        {
            "name": "ticket_purchase",
            "condition": "segment in [L2, L3, L4]",
            "terminal_states": ["Ticket_Transfer", "Ticket_Refund", "Ticket_Admission"],
            "events": [],
        },
        {
            "name": "membership_normal_activation",
            "condition": "segment in [L3, L4]",
            "conversion_rate": 0.80,
            "events": [],
        },
        {
            "name": "membership_late_activation",
            "condition": "segment in [L3, L4]",
            "conversion_rate": 0.10,
            "events": [],
        },
        {
            "name": "membership_expired",
            "condition": "segment in [L3, L4]",
            "conversion_rate": 0.10,
            "events": [],
        },
    ]

    # 6. constraints（骨架）
    data["constraints"] = [
        {
            "type": "temporal_order",
            "description": "Registration 在 Login 前",
        },
        {
            "type": "temporal_order",
            "description": "Login 在 Purchase 前",
        },
        {
            "type": "field_consistency",
            "description": "票数明细数量等于订单票数",
        },
        {
            "type": "field_consistency",
            "description": "明细金额之和不超过订单金额",
        },
        {
            "type": "business_rule",
            "description": "Admission 后不可 Refund",
        },
        {
            "type": "business_rule",
            "description": "Voucher 场景严格互斥",
        },
    ]

    # 7. enums
    data["enums"] = {
        "voucher_scenario": [
            {
                "value": "项目订单",
                "applicable_events": ["Product_Order_Payment"],
            },
            {
                "value": "商品订单",
                "applicable_events": ["Merchandise_Order_Payment"],
            },
        ],
        "voucher_get_type": ["直发", "扫码", "手动领取"],
        "business_unit": [
            "M+博物馆(M)",
            "香港故宫文化博物馆(S)",
            "西九演艺(P)",
            "西九文化区(D)",
        ],
    }

    # 8. failure_rate
    data["failure_rate"] = 0.05

    # 9. fixed_accounts（Table 10，UAT-X01 ~ UAT-X08）
    data["fixed_accounts"] = [
        {
            "id": "UAT-X01",
            "region": "mainland",
            "segment": "L4",
            "crm_master_id": "CRM-UAT-X01",
            "email": "x01@test.westk.hk",
            "mobile": "+86 138-000-0001",
            "unionid": "wxu_x01",
            "cookie_id": "ck_x01_a",
            "split_identity": False,
        },
        {
            "id": "UAT-X02",
            "region": "mainland",
            "segment": "L2",
            "mobile": "+86 138-0000-0002",
            "unionid": "wxu_x02",
            "cookie_id": "ck_x02_a",
            "split_identity": False,
        },
        {
            "id": "UAT-X03",
            "region": "hongkong",
            "segment": "L3",
            "crm_master_id": "CRM-UAT-X03",
            "email": "x03@test.westk.hk",
            "unionid": "wxu_x03",
            "cookie_id": "ck_x03_a",
            "split_identity": False,
        },
        {
            "id": "UAT-X04",
            "region": "hongkong",
            "segment": "L2",
            "email": "x04@test.westk.hk",
            "mobile": "+852 9000-0004",
            "unionid": "wxu_x04",
            "cookie_id": "ck_x04_a",
            "split_identity": False,
        },
        {
            "id": "UAT-X05",
            "region": "overseas",
            "segment": "L1",
            "email": "x05@test.westk.hk",
            "unionid": "wxu_x05",
            "cookie_id": "ck_x05_a",
            "split_identity": False,
        },
        {
            "id": "UAT-X06",
            "region": "mainland",
            "segment": "L4",
            "crm_master_id": "CRM-UAT-X06",
            "email": "x06_a@test.westk.hk",
            "mobile": "+86 138-0000-0006",
            "unionid": "wxu_x06",
            "cookie_ids": ["ck_x06_a", "ck_x06_b", "ck_x06_c"],
            "split_identity": False,
        },
        {
            "id": "UAT-X07",
            "region": "mainland",
            "segment": "L2",
            "email": "x07@test.westk.hk",
            "mobile": "+86 138-0000-0007",
            "unionid": "wxu_x07",
            "cookie_id": "ck_x07_a",
            "split_identity": True,
            "split_groups": [
                {"mobile": "+86 138-0000-0007", "unionid": "wxu_x07"},
                {"email": "x07@test.westk.hk", "cookie_id": "ck_x07_a"},
            ],
        },
        {
            "id": "UAT-X08",
            "region": "mainland",
            "segment": "L2",
            "crm_master_id": "CRM-UAT-X08",
            "email": "x08@corp.westk.hk",
            "mobile": "+86 138-0000-0008",
            "unionid": "wxu_x08",
            "cookie_id": "ck_x08_a",
            "split_identity": False,
            "note": "企业用户，CRM Master ID only，无个人会员卡",
        },
    ]

    return data


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def main():
    # 尝试解析 docx（当前版本直接使用硬编码，解析结果仅用于日志）
    tables, paragraphs = try_parse_docx(DOCX_PATH)
    if tables is not None:
        print(f"[info] docx 解析成功，共 {len(tables)} 个表格，{len(paragraphs)} 个段落")
    else:
        print("[info] 使用硬编码数据生成 YAML")

    data = build_hardcoded_data()

    # 确保输出目录存在
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        f.write("# business_logic.yaml\n")
        f.write("# 由 scripts/docx_to_yaml_skeleton.py 自动生成\n")
        f.write("# 来源：UAT_TestDataset_BusinessLogic_v1.docx\n")
        f.write("#\n")
        f.write("# 此文件为骨架，event_sequences[].events 需手动补充。\n")
        f.write("\n")
        yaml.dump(
            data,
            f,
            allow_unicode=True,
            default_flow_style=False,
            sort_keys=False,
        )

    print(f"[done] 已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    main()
