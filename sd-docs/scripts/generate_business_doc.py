"""
Business Document Generator
Generates professionally formatted Word documents following enterprise standards.

Usage:
    python generate_business_doc.py --input input.md --output output.docx
    python generate_business_doc.py --input input.md --output output.docx --title "Document Title"
"""

import argparse
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class BusinessDocumentFormatter:
    """Formats Word documents according to business document standards."""
    
    # Color constants
    COLOR_BLACK = RGBColor(0, 0, 0)
    COLOR_DARK_BLUE = RGBColor(68, 114, 196)
    COLOR_GRAY = RGBColor(128, 128, 128)
    COLOR_LIGHT_GRAY = RGBColor(242, 242, 242)
    COLOR_WHITE = RGBColor(255, 255, 255)
    COLOR_RED = RGBColor(192, 0, 0)
    
    # Document templates based on industry standards
    DOCUMENT_TEMPLATES = {
        "consulting": {
            "font": "Calibri",
            "font_size": 11,
            "heading_color": "#0F4761",
            "table_header_bg": "#156082",
            "table_alt_rows": True,
            "table_alt_bg": "#F2F2F2",
            "margins": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54},
            "line_spacing": 1.15,
            "header_font": "Calibri",
            "table_font": "Calibri",
        },
        "investment_research": {
            "font": "Arial",
            "font_size": 11,
            "heading_color": "#000000",
            "table_header_bg": "#333333",
            "table_alt_rows": True,
            "table_alt_bg": "#F5F5F5",
            "margins": {"top": 2.0, "bottom": 2.0, "left": 2.0, "right": 2.0},
            "line_spacing": 1.15,
            "header_font": "Arial",
            "table_font": "Arial",
        },
        "technical_spec": {
            "font": "Arial",
            "font_size": 11,
            "heading_color": "#000000",
            "table_header_bg": "#4472C4",
            "table_alt_rows": True,
            "table_alt_bg": "#F2F2F2",
            "margins": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 2.54},
            "line_spacing": 1.15,
            "header_font": "Arial",
            "table_font": "Arial",
        },
        "default": {
            "font": "Calibri",
            "font_size": 11,
            "heading_color": "#000000",
            "table_header_bg": "#4472C4",
            "table_alt_rows": True,
            "table_alt_bg": "#F2F2F2",
            "margins": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 2.54},
            "line_spacing": 1.15,
            "header_font": "Calibri",
            "table_font": "Calibri",
        }
    }
    
    def __init__(self, document_title="", subtitle="", version="1.0", 
                 author="", classification="Confidential", template="default"):
        self.doc = Document()
        self.document_title = document_title
        self.subtitle = subtitle
        self.version = version
        self.author = author
        self.classification = classification
        self.template = template
        self.config = self.DOCUMENT_TEMPLATES.get(template, self.DOCUMENT_TEMPLATES["default"])
        self._setup_page()
        self._setup_styles()
    
    def _setup_page(self):
        """Configure page layout and margins based on template."""
        margins = self.config["margins"]
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(margins["top"])
        section.bottom_margin = Cm(margins["bottom"])
        section.left_margin = Cm(margins["left"])
        section.right_margin = Cm(margins["right"])
        
    def _setup_styles(self):
        """Configure document styles based on template."""
        font_name = self.config["font"]
        font_size = self.config["font_size"]
        heading_color = self.config["heading_color"]
        line_spacing = self.config["line_spacing"]
        
        # Convert hex color to RGBColor
        if heading_color.startswith('#'):
            heading_rgb = RGBColor(
                int(heading_color[1:3], 16),
                int(heading_color[3:5], 16),
                int(heading_color[5:7], 16)
            )
        else:
            heading_rgb = self.COLOR_BLACK
        
        # Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.color.rgb = self.COLOR_BLACK
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.space_before = Pt(0)
        
        # Heading styles
        heading_configs = [
            ('Heading 1', 16, True, Pt(18), Pt(12)),
            ('Heading 2', 14, True, Pt(14), Pt(8)),
            ('Heading 3', 12, True, Pt(10), Pt(6)),
            ('Heading 4', 11, True, Pt(8), Pt(4)),
        ]
        
        for style_name, size, bold, space_before, space_after in heading_configs:
            try:
                style = self.doc.styles[style_name]
            except KeyError:
                style = self.doc.styles.add_style(style_name, 1)  # WD_STYLE_TYPE.PARAGRAPH
            font = style.font
            font.name = font_name
            font.size = Pt(size)
            font.bold = bold
            font.color.rgb = heading_rgb
            style.paragraph_format.space_before = space_before
            style.paragraph_format.space_after = space_after
            style.paragraph_format.line_spacing = line_spacing
            style.paragraph_format.keep_with_next = True
    
    def add_cover_page(self, company_logo_path=None):
        """Add professional cover page."""
        font_name = self.config["font"]
        
        # Company logo (placeholder text if no image)
        if company_logo_path:
            self.doc.add_picture(company_logo_path, width=Inches(2.0))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            logo_para = self.doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run("[Company Logo]")
            logo_run.font.size = Pt(12)
            logo_run.font.color.rgb = self.COLOR_GRAY
            logo_run.font.name = font_name
        
        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(self.document_title)
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.name = font_name
        title_run.font.color.rgb = self.COLOR_BLACK
        
        # Subtitle
        if self.subtitle:
            self.doc.add_paragraph()
            subtitle_para = self.doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(self.subtitle)
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.name = font_name
            subtitle_run.font.color.rgb = self.COLOR_BLACK
        
        # Spacing
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Version and date
        meta_para = self.doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_text = f"Version: {self.version}\nDate: {datetime.now().strftime('%Y-%m-%d')}"
        if self.author:
            meta_text += f"\nAuthor: {self.author}"
        meta_run = meta_para.add_run(meta_text)
        meta_run.font.size = Pt(11)
        meta_run.font.name = font_name
        meta_run.font.color.rgb = self.COLOR_GRAY
        
        # Push classification to bottom
        for _ in range(6):
            self.doc.add_paragraph()
        
        # Classification
        class_para = self.doc.add_paragraph()
        class_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        class_run = class_para.add_run(self.classification.upper())
        class_run.font.size = Pt(10)
        class_run.font.bold = True
        class_run.font.name = font_name
        class_run.font.color.rgb = self.COLOR_RED
        
        # Page break
        self.doc.add_page_break()
    
    def add_document_control(self, metadata=None):
        """Add document control table."""
        font_name = self.config["font"]
        
        if metadata is None:
            metadata = {
                'Document Name': self.document_title,
                'Version': self.version,
                'Date': datetime.now().strftime('%Y-%m-%d'),
                'Author': self.author or '[Author Name]',
                'Reviewer': '[Reviewer Name]',
                'Approver': '[Approver Name]',
                'Status': 'Draft',
                'Classification': self.classification,
            }
        
        # Title
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run('Document Control')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.name = font_name
        title_para.paragraph_format.space_after = Pt(12)
        
        # Table
        table = self.doc.add_table(rows=len(metadata), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for idx, (key, value) in enumerate(metadata.items()):
            row = table.rows[idx]
            
            # Key cell
            key_cell = row.cells[0]
            key_cell.text = key
            key_para = key_cell.paragraphs[0]
            key_run = key_para.runs[0]
            key_run.font.bold = True
            key_run.font.size = Pt(11)
            key_run.font.name = font_name
            key_cell.width = Cm(5)
            
            # Value cell
            value_cell = row.cells[1]
            value_cell.text = str(value)
            value_para = value_cell.paragraphs[0]
            value_run = value_para.runs[0]
            value_run.font.size = Pt(11)
            value_run.font.name = font_name
            value_cell.width = Cm(10)
            
            # Cell shading for header column
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F2F2F2')
            key_cell._tc.get_or_add_tcPr().append(shading)
        
        self.doc.add_paragraph()
        self.doc.add_page_break()
    
    def add_toc_placeholder(self):
        """Add TOC placeholder with instructions."""
        font_name = self.config["font"]
        
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run('Table of Contents')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.name = font_name
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(18)
        
        instruction = self.doc.add_paragraph()
        instruction_run = instruction.add_run(
            '[Update Table of Contents: References → Table of Contents → Update Table]'
        )
        instruction_run.font.size = Pt(10)
        instruction_run.font.italic = True
        instruction_run.font.color.rgb = self.COLOR_GRAY
        instruction_run.font.name = font_name
        instruction.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()
    
    def add_header_footer(self):
        """Add header and footer to all sections."""
        font_name = self.config["font"]
        header_font = self.config["header_font"]
        
        for section in self.doc.sections:
            # Header
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = self.document_title
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in header_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = self.COLOR_GRAY
                run.font.name = header_font
            
            # Add bottom border to header
            pPr = header_para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '808080')
            pBdr.append(bottom)
            pPr.append(pBdr)
            
            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            
            # Left: Classification
            footer_para.text = self.classification
            footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in footer_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = self.COLOR_RED
                run.font.name = font_name
            
            # Right: Page number
            # Add tab for right alignment
            run = footer_para.add_run('\t')
            run = footer_para.add_run('Page ')
            run.font.size = Pt(9)
            run.font.color.rgb = self.COLOR_GRAY
            run.font.name = font_name
            
            # Add page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            run = footer_para.add_run()
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run.font.size = Pt(9)
            run.font.color.rgb = self.COLOR_GRAY
            run.font.name = font_name
    
    def format_table(self, table, header_row=True):
        """Apply professional formatting to a table based on template."""
        table_font = self.config["table_font"]
        table_header_bg = self.config["table_header_bg"]
        table_alt_rows = self.config["table_alt_rows"]
        table_alt_bg = self.config["table_alt_bg"]
        
        if header_row and len(table.rows) > 0:
            # Format header row
            header_row_obj = table.rows[0]
            for cell in header_row_obj.cells:
                # Background color
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), table_header_bg.lstrip('#'))
                cell._tc.get_or_add_tcPr().append(shading)
                
                # Text formatting
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = self.COLOR_WHITE
                        run.font.size = Pt(10.5)
                        run.font.name = table_font
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Format all cells
        for row_idx, row in enumerate(table.rows):
            if header_row and row_idx == 0:
                continue
            
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10.5)
                        run.font.name = table_font
                    
                    # Alternating row colors
                    if table_alt_rows and row_idx % 2 == 0:
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), table_alt_bg.lstrip('#'))
                        cell._tc.get_or_add_tcPr().append(shading)
                
                # Cell margins
                tcPr = cell._tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for edge in ['top', 'left', 'bottom', 'right']:
                    edge_el = OxmlElement(f'w:{edge}')
                    edge_el.set(qn('w:w'), '80')
                    edge_el.set(qn('w:type'), 'dxa')
                    tcMar.append(edge_el)
                tcPr.append(tcMar)
    
    def add_heading(self, text, level=1):
        """Add formatted heading."""
        style_name = f'Heading {level}'
        para = self.doc.add_paragraph(text, style=style_name)
        return para
    
    def add_paragraph(self, text, bold=False, italic=False, 
                     color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """Add formatted paragraph."""
        font_name = self.config["font"]
        font_size = self.config["font_size"]
        
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = font_name
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        para.alignment = alignment
        return para
    
    def save(self, output_path):
        """Save document."""
        self.doc.save(output_path)
        print(f"Document saved: {output_path}")


def generate_from_markdown(md_content, output_path, **kwargs):
    """Generate formatted Word document from Markdown content."""
    formatter = BusinessDocumentFormatter(**kwargs)
    
    # Add front matter
    formatter.add_cover_page()
    formatter.add_document_control()
    formatter.add_toc_placeholder()
    
    # Parse markdown and add content
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip YAML front matter
        if line == '---':
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue
        
        # Skip table of contents section
        if line.lower() in ['## table of contents', '## 目录']:
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('#'):
                i += 1
            continue
        
        # Headings
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            formatter.add_heading(text, level=min(level, 4))
            i += 1
            continue
        
        # Tables (simple parsing)
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            # Parse table
            rows = []
            while i < len(lines) and '|' in lines[i]:
                cells = [cell.strip() for cell in lines[i].split('|')]
                cells = [c for c in cells if c]  # Remove empty
                if cells and not all('-' in c for c in cells):
                    rows.append(cells)
                i += 1
            
            if rows:
                num_cols = max(len(row) for row in rows)
                table = formatter.doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(rows):
                    row = table.rows[row_idx]
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            row.cells[col_idx].text = cell_text
                
                formatter.format_table(table, header_row=True)
            continue
        
        # Regular paragraph
        if line:
            formatter.add_paragraph(line)
        
        i += 1
    
    # Add header/footer
    formatter.add_header_footer()
    
    # Save
    formatter.save(output_path)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate business document')
    parser.add_argument('--input', '-i', help='Input markdown file')
    parser.add_argument('--output', '-o', required=True, help='Output docx file')
    parser.add_argument('--title', '-t', default='Business Document', help='Document title')
    parser.add_argument('--subtitle', '-s', default='', help='Document subtitle')
    parser.add_argument('--version', '-v', default='1.0', help='Document version')
    parser.add_argument('--author', '-a', default='', help='Document author')
    parser.add_argument('--classification', '-c', default='Confidential', 
                       help='Document classification')
    parser.add_argument('--template', '-tp', default='default', 
                       choices=['consulting', 'investment_research', 'technical_spec', 'default'],
                       help='Document template style')
    
    args = parser.parse_args()
    
    kwargs = {
        'document_title': args.title,
        'subtitle': args.subtitle,
        'version': args.version,
        'author': args.author,
        'classification': args.classification,
        'template': args.template,
    }
    
    if args.input:
        with open(args.input, 'r', encoding='utf-8') as f:
            md_content = f.read()
        generate_from_markdown(md_content, args.output, **kwargs)
    else:
        # Generate blank template
        formatter = BusinessDocumentFormatter(**kwargs)
        formatter.add_cover_page()
        formatter.add_document_control()
        formatter.add_toc_placeholder()
        formatter.add_heading('1. Introduction', level=1)
        formatter.add_paragraph('This is a sample paragraph. Replace with your content.')
        formatter.add_heading('1.1 Background', level=2)
        formatter.add_paragraph('Subsection content goes here.')
        formatter.add_header_footer()
        formatter.save(args.output)
