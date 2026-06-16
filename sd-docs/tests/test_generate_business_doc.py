"""Tests for generate_business_doc.py."""
from docx import Document

from generate_business_doc import BusinessDocumentFormatter, generate_from_markdown


def test_formatter_creates_document(tmp_path):
    formatter = BusinessDocumentFormatter(
        document_title="Test Doc",
        subtitle="Subtitle",
        version="1.0",
        author="Tester",
        template="default",
    )
    formatter.add_cover_page()
    formatter.add_document_control()
    formatter.add_toc_placeholder()
    formatter.add_heading("1. Introduction", level=1)
    formatter.add_paragraph("This is content.")
    formatter.add_header_footer()

    out = tmp_path / "test.docx"
    formatter.save(str(out))
    assert out.exists()

    doc = Document(str(out))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Test Doc" in texts
    assert "Introduction" in texts


def test_generate_from_markdown(tmp_path):
    md = """# Business Document

This is the introduction.

## Scope

| Item | Value |
|------|-------|
| A    | 1     |
| B    | 2     |
"""
    out = tmp_path / "out.docx"
    generate_from_markdown(md, str(out), document_title="MD Doc")
    assert out.exists()

    doc = Document(str(out))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Business Document" in texts
    assert "Scope" in texts


def test_generate_from_markdown_skips_front_matter(tmp_path):
    md = """---
title: Test
---
# Content

Hello.
"""
    out = tmp_path / "out.docx"
    generate_from_markdown(md, str(out), document_title="Doc")
    doc = Document(str(out))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "title: Test" not in texts
    assert "Content" in texts


def test_blank_template_generation(tmp_path):
    """Test CLI path that generates blank template when no input provided."""
    formatter = BusinessDocumentFormatter(document_title="Blank")
    formatter.add_cover_page()
    formatter.add_document_control()
    formatter.add_toc_placeholder()
    formatter.add_heading('1. Introduction', level=1)
    formatter.add_paragraph('This is a sample paragraph. Replace with your content.')
    formatter.add_header_footer()
    out = tmp_path / "blank.docx"
    formatter.save(str(out))
    assert out.exists()


def test_different_templates():
    for template in ["consulting", "investment_research", "technical_spec", "default"]:
        formatter = BusinessDocumentFormatter(template=template)
        assert formatter.config["font"] is not None
