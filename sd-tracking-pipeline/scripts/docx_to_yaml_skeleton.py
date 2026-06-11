"""
docx_to_yaml_skeleton.py
========================
从业务需求文档（docx）提取表格数据，生成 business_logic.yaml 骨架。

运行方式：
    cd tracking-setup-e2e
    python3 scripts/docx_to_yaml_skeleton.py \
        --input ./references/business_logic.docx \
        --output ./rules/business_logic.yaml \
        --project myproject

依赖（可选）：
    pip install python-docx pyyaml

如果 docx 解析失败，脚本会使用通用模板数据生成 YAML，
用户需手动补充 event_sequences 和 fixed_accounts。
"""

import argparse
import os
import sys
import yaml


# ---------------------------------------------------------------------------
# 通用模板数据
# ---------------------------------------------------------------------------

def build_template_data(project: str) -> dict:
    """生成通用 business_logic.yaml 模板，不含任何客户特定业务逻辑。"""
    data = {}

    # 1. meta
    data["meta"] = {
        "project": project,
        "version": "1.0",
        "platforms": ["MP", "Web"],
    }

    # 2. region_distribution（全局）
    data["region_distribution"] = {
        "mainland": 0.60,
        "hongkong": 0.30,
        "overseas": 0.10,
    }

    # 3. user_segments（示例分层，可按需调整）
    data["user_segments"] = {
        "L0": {
            "ratio": 0.60,
            "identities": ["unionid", "cookie_id"],
            "has_registration": False,
            "has_purchase": False,
            "has_membership": False,
            "membership_activated": False,
        },
        "L1": {
            "ratio": 0.20,
            "identities": ["mobile_or_email", "unionid", "cookie_id"],
            "has_registration": True,
            "has_purchase": False,
            "has_membership": False,
            "membership_activated": False,
        },
        "L2": {
            "ratio": 0.10,
            "identities": ["mobile_or_email", "unionid", "cookie_id"],
            "has_registration": True,
            "has_purchase": True,
            "has_membership": False,
            "membership_activated": False,
        },
        "L3": {
            "ratio": 0.03,
            "identities": ["mobile_or_email", "unionid", "cookie_id"],
            "has_registration": True,
            "has_purchase": True,
            "has_membership": True,
            "membership_activated": False,
        },
        "L4": {
            "ratio": 0.07,
            "identities": ["mobile_or_email", "unionid", "cookie_id"],
            "has_registration": True,
            "has_purchase": True,
            "has_membership": True,
            "membership_activated": True,
        },
    }

    # 4. identity_priority
    data["identity_priority"] = {
        "crm_master_id": {
            "priority": 0,
            "display": "CRM Master Customer ID",
            "sa_key": "$identity_login_id",
        },
        "email": {
            "priority": 1,
            "display": "Email",
            "sa_key": "$identity_email",
        },
        "mobile": {
            "priority": 1,
            "display": "Mobile",
            "sa_key": "$identity_mobile",
        },
        "unionid": {
            "priority": 2,
            "display": "WeChat UnionID",
            "sa_key": "$identity_unionid",
        },
        "cookie_id": {
            "priority": 2,
            "display": "CookieID",
            "sa_key": "$identity_anonymous_id",
        },
    }

    # 5. event_sequences（骨架，需手动补充 events 列表）
    data["event_sequences"] = [
        {
            "name": "user_lifecycle",
            "condition": None,
            "events": [],
        },
        {
            "name": "purchase_flow",
            "condition": "segment in [L2, L3, L4]",
            "terminal_states": [],
            "events": [],
        },
        {
            "name": "membership_activation",
            "condition": "segment in [L3, L4]",
            "conversion_rate": 0.80,
            "events": [],
        },
    ]

    # 6. constraints（骨架，需根据实际业务规则调整）
    data["constraints"] = [
        {
            "type": "temporal_order",
            "description": "Registration 在 Login 前",
        },
        {
            "type": "temporal_order",
            "description": "Login 在 Purchase 前",
        },
        {
            "type": "field_consistency",
            "description": "明细数量等于订单数量",
        },
        {
            "type": "field_consistency",
            "description": "明细金额之和不超过订单金额",
        },
        {
            "type": "business_rule",
            "description": "Admission 后不可 Refund",
        },
    ]

    # 7. enums（骨架，需根据实际业务枚举值填充）
    data["enums"] = {
        "voucher_scenario": [
            {"value": "项目订单", "applicable_events": ["Product_Order_Payment"]},
            {"value": "商品订单", "applicable_events": ["Merchandise_Order_Payment"]},
        ],
        "voucher_get_type": ["直发", "扫码", "手动领取"],
    }

    # 8. failure_rate
    data["failure_rate"] = 0.05

    # 9. fixed_accounts（骨架，需根据 UAT 需求手动补充）
    data["fixed_accounts"] = [
        {
            "id": "UAT-001",
            "region": "mainland",
            "segment": "L4",
            "crm_master_id": "CRM-UAT-001",
            "email": "uat001@example.com",
            "mobile": "+86 138-0000-0001",
            "unionid": "wxu_uat001",
            "cookie_id": "ck_uat001_a",
            "split_identity": False,
        },
        {
            "id": "UAT-002",
            "region": "mainland",
            "segment": "L2",
            "mobile": "+86 138-0000-0002",
            "unionid": "wxu_uat002",
            "cookie_id": "ck_uat002_a",
            "split_identity": False,
        },
        {
            "id": "UAT-003",
            "region": "hongkong",
            "segment": "L3",
            "crm_master_id": "CRM-UAT-003",
            "email": "uat003@example.com",
            "unionid": "wxu_uat003",
            "cookie_id": "ck_uat003_a",
            "split_identity": False,
        },
    ]

    return data


# ---------------------------------------------------------------------------
# docx 解析（可选）
# ---------------------------------------------------------------------------

def try_parse_docx(path: str):
    """尝试用 python-docx 读取文档，返回 (tables, paragraphs) 或 (None, None)。"""
    try:
        from docx import Document
        doc = Document(path)
        return doc.tables, doc.paragraphs
    except Exception as e:
        print(f"[warn] docx 解析失败，使用通用模板：{e}", file=sys.stderr)
        return None, None


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="从业务需求文档生成 business_logic.yaml 骨架"
    )
    parser.add_argument("--input", default="", help="输入 docx 文件路径（可选，不提供则使用通用模板）")
    parser.add_argument("--output", required=True, help="输出 YAML 文件路径")
    parser.add_argument("--project", default="myproject", help="项目标识名（默认: myproject）")
    args = parser.parse_args()

    # 尝试解析 docx
    if args.input and os.path.exists(args.input):
        tables, paragraphs = try_parse_docx(args.input)
        if tables is not None:
            print(f"[info] docx 解析成功，共 {len(tables)} 个表格，{len(paragraphs)} 个段落")
            print("[info] 注意：docx 解析结果未自动映射到 YAML 结构，当前版本使用通用模板")
        else:
            print("[info] 使用通用模板生成 YAML")
    else:
        if args.input:
            print(f"[warn] 找不到输入文件: {args.input}，使用通用模板")
        else:
            print("[info] 未提供输入文件，使用通用模板")

    data = build_template_data(args.project)

    # 确保输出目录存在
    output_dir = os.path.dirname(os.path.abspath(args.output))
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    with open(args.output, "w", encoding="utf-8") as f:
        f.write("# business_logic.yaml\n")
        f.write("# 由 scripts/docx_to_yaml_skeleton.py 自动生成\n")
        if args.input:
            f.write(f"# 来源：{os.path.basename(args.input)}\n")
        f.write("#\n")
        f.write("# 此文件为通用模板，以下部分需根据实际业务需求手动补充：\n")
        f.write("#   1. event_sequences[].events — 每个序列的具体事件列表\n")
        f.write("#   2. constraints — 业务约束规则\n")
        f.write("#   3. enums — 业务枚举值\n")
        f.write("#   4. fixed_accounts — UAT 测试账号（替换示例数据）\n")
        f.write("#   5. property_enums — 属性级别的枚举覆盖（可选）\n")
        f.write("#   6. preset_events — 小程序预置事件配置（如适用）\n")
        f.write("\n")
        yaml.dump(
            data,
            f,
            allow_unicode=True,
            default_flow_style=False,
            sort_keys=False,
        )

    print(f"[done] 已生成：{args.output}")
    print(f"[hint] 请根据实际业务需求编辑该文件，然后运行 yaml_validator.py 验证")


if __name__ == "__main__":
    main()
